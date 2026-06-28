"""
resume_parser.py
-----------------
Layer 1 of the pipeline: turn a raw resume file (PDF / DOCX / TXT) into clean
text, then split that text into labeled sections (Education, Experience,
Skills, Projects, Certifications).

Why section detection matters: once we know "this paragraph is the Education
section", every downstream extractor (skills, dates, degree level) can look
in the right place instead of guessing across the whole document. This single
step is responsible for a large chunk of the accuracy gain over naive
whole-document parsing.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

# Canonical section name -> regex patterns that mark the start of that section.
# Order matters: more specific patterns should come first within a list.
SECTION_HEADERS: Dict[str, List[str]] = {
    "summary": [r"professional\s+summary", r"summary", r"profile", r"objective"],
    "skills": [r"technical\s+skills", r"skills", r"core\s+competencies", r"technologies"],
    "experience": [
        r"work\s+experience",
        r"professional\s+experience",
        r"experience",
        r"employment\s+history",
    ],
    "education": [r"education", r"academic\s+background", r"qualifications"],
    "projects": [r"projects", r"personal\s+projects", r"academic\s+projects"],
    "certifications": [r"certifications?", r"licenses?", r"courses"],
    "contact": [r"contact", r"personal\s+information"],
}


@dataclass
class ParsedResume:
    """Container for a resume after extraction + section splitting."""

    source_path: str
    raw_text: str
    sections: Dict[str, str] = field(default_factory=dict)
    parse_warnings: List[str] = field(default_factory=list)

    def section(self, name: str) -> str:
        """Return the text of a section, or '' if it wasn't found."""
        return self.sections.get(name, "")


def extract_text_from_pdf(path: Path) -> str:
    """Extract text from a PDF, page by page, preserving line breaks."""
    import pdfplumber

    chunks: List[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            chunks.append(text)
    return "\n".join(chunks)


def extract_text_from_docx(path: Path) -> str:
    """Extract text from a .docx file, paragraph by paragraph."""
    import docx

    document = docx.Document(str(path))
    lines = [p.text for p in document.paragraphs]
    # Tables (some resumes use a table layout for contact info / skills)
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(lines)


def extract_text_from_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_text(path: Path) -> str:
    """Dispatch to the right extractor based on file extension."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    if ext == ".docx":
        return extract_text_from_docx(path)
    if ext == ".txt":
        return extract_text_from_txt(path)
    raise ValueError(
        f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
    )


def _build_header_regex() -> re.Pattern:
    """Combine all section header patterns into one alternation regex.

    A line is treated as a header if, after stripping punctuation, it matches
    one of the known header patterns AND is short (headers are rarely full
    sentences) AND doesn't end with typical sentence punctuation.
    """
    all_patterns = []
    for patterns in SECTION_HEADERS.values():
        all_patterns.extend(patterns)
    combined = "|".join(all_patterns)
    return re.compile(rf"^\s*({combined})\s*[:.\-]?\s*$", re.IGNORECASE)


_HEADER_RE = _build_header_regex()


def _classify_header(line: str) -> str | None:
    """Return the canonical section name if `line` looks like a header."""
    cleaned = line.strip()
    if not cleaned or len(cleaned) > 40:
        return None
    if not _HEADER_RE.match(cleaned):
        return None
    lower = cleaned.lower()
    for canonical, patterns in SECTION_HEADERS.items():
        for pat in patterns:
            if re.fullmatch(pat, lower.strip(":.- ")):
                return canonical
    return None


def split_into_sections(raw_text: str) -> Dict[str, str]:
    """Split resume text into sections keyed by canonical section name.

    Any text appearing before the first recognized header is kept under the
    'header' key (this is usually name / contact info / a one-line title).
    """
    lines = raw_text.splitlines()
    sections: Dict[str, List[str]] = {"header": []}
    current = "header"

    for line in lines:
        section_name = _classify_header(line)
        if section_name:
            current = section_name
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)

    return {name: "\n".join(lines).strip() for name, lines in sections.items()}


def parse_resume(path: str | Path) -> ParsedResume:
    """Full Layer-1 pipeline: read file -> extract text -> split sections."""
    path = Path(path)
    warnings: List[str] = []

    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {path}")
    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{path.suffix}'. "
            f"Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    raw_text = extract_text(path)

    if not raw_text.strip():
        warnings.append(
            "No extractable text found. The file may be a scanned/image-based "
            "PDF — OCR (e.g. pytesseract) would be needed as a fallback."
        )

    sections = split_into_sections(raw_text)

    if len(sections) <= 1:
        warnings.append(
            "No section headers were detected — the whole document was kept "
            "under 'header'. Downstream extractors will fall back to "
            "searching the full text instead of per-section text."
        )

    return ParsedResume(
        source_path=str(path),
        raw_text=raw_text,
        sections=sections,
        parse_warnings=warnings,
    )
